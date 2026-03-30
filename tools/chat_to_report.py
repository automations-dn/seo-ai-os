#!/usr/bin/env python3
"""
Dare Network Output Formatter
Converts approved Mastermind Agent output (.md) into styled .docx.
NO recalculations. NO hardcoded logic. What you see is what you get.
"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm

NAVY = RGBColor(0x1B, 0x3A, 0x6B)
ORANGE = RGBColor(0xE8, 0x67, 0x1A)
DARK = RGBColor(0x33, 0x33, 0x33)

def _add_formatted_text(paragraph, text):
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        r = paragraph.add_run(part)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        if i % 2 == 1:
            r.bold = True

def convert_md_to_docx(input_md, output_docx):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    with open(input_md, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False

    for line in lines:
        line = line.rstrip()

        # Track code blocks
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            r.font.size = Pt(10)
            r.font.name = "Consolas"
            r.font.color.rgb = DARK
            continue

        if not line.strip():
            doc.add_paragraph()
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = NAVY
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = NAVY
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run(line[4:])
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = ORANGE
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            continue

        # Lists
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, line.strip()[2:])
            continue

        p = doc.add_paragraph()
        _add_formatted_text(p, line)

    doc.save(output_docx)
    print(f"[SUCCESS] Report styles applied. Saved to: {output_docx}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Convert Agent Markdown to Dare Network DOCX")
    parser.add_argument("--input", required=True, help="Input markdown file (.md)")
    parser.add_argument("--output", required=True, help="Output word document (.docx)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[ERROR] Input file not found: {input_path}")
        sys.exit(1)

    convert_md_to_docx(args.input, args.output)
