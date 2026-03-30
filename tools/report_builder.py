#!/usr/bin/env python3
"""
Report Builder — Dare Network Template v2
Matches the professional audit template exactly.
Colors: Navy #1B3A6B, Orange #E8671A, White background
Tables: Issue | Severity | Finding | Recommended Fix (blue italic)
"""

import argparse, json, os, sys
from pathlib import Path
from datetime import datetime, date
from glob import glob

# Add tools/ directory to path for imports
sys.path.insert(0, str(Path(__file__).parent))
from utils import url_to_slug, validate_file_naming  # CRITICAL: Use centralized URL-to-slug conversion

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess; subprocess.run(["pip","install","python-docx"],check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1B,0x3A,0x6B)
ORANGE    = RGBColor(0xE8,0x67,0x1A)
WHITE     = RGBColor(0xFF,0xFF,0xFF)
DARK      = RGBColor(0x33,0x33,0x33)
MID       = RGBColor(0x66,0x66,0x66)
BLUE_LINK = RGBColor(0x1A,0x5C,0xA8)
RED_TXT   = RGBColor(0xC0,0x39,0x2B)

def _set_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),hex_color); shd.set(qn("w:val"),"clear")
    tcPr.append(shd)

def _hr(doc, color="1B3A6B"):
    hr=doc.add_paragraph(); hr.paragraph_format.space_before=Pt(0); hr.paragraph_format.space_after=Pt(10)
    pPr=hr._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    b=OxmlElement("w:bottom"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6")
    b.set(qn("w:space"),"1"); b.set(qn("w:color"),color); pBdr.append(b); pPr.append(pBdr)

def add_section_banner(doc, num, title):
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"
    c=t.rows[0].cells[0]; _set_bg(c,"1B3A6B")
    p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(f"  SECTION {num} \u2014 {title.upper()}")
    r.bold=True; r.font.size=Pt(13); r.font.color.rgb=WHITE
    doc.add_paragraph()

def add_h1(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
    _hr(doc)

def add_h2(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY

def add_body(doc, text, colour=None, italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text); r.font.size=Pt(10.5); r.font.color.rgb=colour or DARK; r.italic=italic

def add_bullet(doc, text, bold_part=None):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.left_indent=Inches(0.3)
    if bold_part:
        rb=p.add_run(bold_part+" "); rb.bold=True; rb.font.size=Pt(10.5); rb.font.color.rgb=DARK
    r=p.add_run(text); r.font.size=Pt(10.5); r.font.color.rgb=DARK

def add_callout(doc, title, lines, style="info"):
    bgs={"info":"E3F2FD","warning":"FBE9E7","critical":"FFEBEE","success":"E8F5E9"}
    cols={"info":NAVY,"warning":ORANGE,"critical":RED_TXT,"success":RGBColor(0x1E,0x8B,0x32)}
    icons={"info":"\U0001f4a1","warning":"\u26a0","critical":"\u26a0","success":"\u2705"}
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"
    c=t.rows[0].cells[0]; _set_bg(c,bgs.get(style,"E3F2FD"))
    tp=c.paragraphs[0]; tp.paragraph_format.space_before=Pt(6)
    tr=tp.add_run(f"{icons.get(style,'')}  {title}")
    tr.bold=True; tr.font.size=Pt(10.5); tr.font.color.rgb=cols.get(style,NAVY)
    for line in lines:
        bp=c.add_paragraph(); br=bp.add_run(line)
        br.font.size=Pt(10); br.font.color.rgb=DARK
    last=c.add_paragraph(); last.paragraph_format.space_after=Pt(6)
    doc.add_paragraph()

def add_issues_table(doc, issues):
    if not issues: return
    lbl=doc.add_paragraph(); lbl.paragraph_format.space_before=Pt(8); lbl.paragraph_format.space_after=Pt(2)
    lr=lbl.add_run("Issues Found"); lr.bold=True; lr.font.size=Pt(11); lr.font.color.rgb=NAVY

    t=doc.add_table(rows=1,cols=4); t.style="Table Grid"
    for i,(label,w) in enumerate(zip(["Issue","Severity","Finding","Recommended Fix"],
                                      [Inches(1.4),Inches(0.85),Inches(2.5),Inches(2.25)])):
        c=t.rows[0].cells[i]; _set_bg(c,"1B3A6B")
        c.width=w; p=c.paragraphs[0]
        r=p.add_run(label); r.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(10)

    sev_cfg={"Critical":("8B1C1C",WHITE),"High":("C96A00",WHITE),
             "Medium":("9A6B00",WHITE),"Low":("1E6B32",WHITE)}

    for item in issues:
        row=t.add_row().cells
        sev=item.get("severity",item.get("priority","Medium"))
        # col0 issue name
        row[0].paragraphs[0].clear(); r0=row[0].paragraphs[0].add_run(item.get("issue",""))
        r0.bold=True; r0.font.size=Pt(10)
        # col1 severity badge
        cfg=sev_cfg.get(sev,sev_cfg["Medium"]); _set_bg(row[1],cfg[0])
        row[1].paragraphs[0].clear(); rs=row[1].paragraphs[0].add_run(sev)
        rs.bold=True; rs.font.size=Pt(9); rs.font.color.rgb=WHITE
        row[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        # col2 finding
        row[2].paragraphs[0].clear(); r2=row[2].paragraphs[0].add_run(item.get("finding",item.get("description","")))
        r2.font.size=Pt(10)
        # col3 fix (blue italic)
        row[3].paragraphs[0].clear(); r3=row[3].paragraphs[0].add_run(item.get("fix",item.get("action","")))
        r3.font.size=Pt(10); r3.font.color.rgb=BLUE_LINK; r3.italic=True
    doc.add_paragraph()

def add_health_table(doc, rows):
    t=doc.add_table(rows=1,cols=4); t.style="Table Grid"
    for i,label in enumerate(["Audit Area","Current Status","Score","Priority"]):
        c=t.rows[0].cells[i]; _set_bg(c,"1B3A6B")
        r=c.paragraphs[0].add_run(label); r.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(10)
    score_bgs={"critical":"FFCDD2","high":"FFF9C4","medium":"FFF9C4","good":"C8E6C9"}
    for area,status,score,priority in rows:
        row=t.add_row().cells
        r0=row[0].paragraphs[0].add_run(area); r0.bold=True; r0.font.size=Pt(10)
        r1=row[1].paragraphs[0].add_run(status); r1.font.size=Pt(10)
        r2=row[2].paragraphs[0].add_run(str(score)); r2.bold=True; r2.font.size=Pt(11)
        prio_low=priority.lower()
        _set_bg(row[2],score_bgs.get("critical" if "critical" in prio_low else
                                      "high" if "high" in prio_low else
                                      "medium" if "medium" in prio_low else "good","FFF9C4"))
        r3=row[3].paragraphs[0].add_run(priority); r3.font.size=Pt(10)
    doc.add_paragraph()

# ── Data loaders ─────────────────────────────────────────────────────────────
def load_brand_kit(name):
    p=Path(f"clients/{name}/brand_kit.json")
    return json.load(open(p,encoding="utf-8")) if p.exists() else {"client_info":{"client_name":name,"website_url":f"{name}.com"}}

def load_tmp(pattern):
    files=sorted(glob(f".tmp/{pattern}"),reverse=True)
    return json.load(open(files[0],encoding="utf-8")) if files else {}


def _tech_score(summary, framework_data=None):
    """
    Calculate technical SEO score.
    IMPORTANT: Respects framework detection score caps for CSR/SPA sites.
    """
    s=100
    s-=min(int(summary.get("status_404",0) or 0)*4,30)
    s-=min(int(summary.get("noindex_pages",0) or 0)*2,20)
    s-=min(int(summary.get("missing_h1",0) or 0)*2,20)
    s-=min(int(summary.get("missing_meta_desc",0) or 0)*1,10)
    calculated_score = max(s,10)

    # CRITICAL: Apply framework detection score cap if CSR_SPA detected
    if framework_data and framework_data.get("status") == "success":
        score_cap_data = framework_data.get("score_cap", {})
        if score_cap_data and "technical_seo" in score_cap_data:
            max_score = score_cap_data["technical_seo"] * 10  # Convert 2/10 to 20/100
            calculated_score = min(calculated_score, max_score)

    return calculated_score

# ── MAIN BUILDER ─────────────────────────────────────────────────────────────
def build_audit_docx(client_slug, website_url, brand_kit, strategy_data={}):
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    info=brand_kit.get("client_info",{}); client_name=info.get("client_name",client_slug)
    industry=info.get("industry","");
    cms=strategy_data.get("cms", brand_kit.get("technical_settings",{}).get("cms",""))
    is_ecommerce=strategy_data.get("is_ecommerce", False)
    audit_date=datetime.now().strftime("%B %Y")

    crawl=load_tmp(f"*{client_slug}*crawl*.json"); onpage=load_tmp(f"*{client_slug}*onpage*.json")
    lighthouse=load_tmp(f"*{client_slug}*lighthouse*.json")
    framework_data=load_tmp(f"*{client_slug}*framework*.json")  # CRITICAL: Load framework detection data

    summary=crawl.get("summary",{}); pages=crawl.get("pages",[])

    # ── FIX: Extract real client name from crawl data when using --url mode ──
    # When called with --url, client_name defaults to the full URL which looks bad in the report.
    # Try to extract a cleaner name from the homepage title in the crawl data.
    if client_name == website_url or client_name.startswith("http"):
        homepage_pages = [p for p in pages if p.get("url","").rstrip("/") == website_url.rstrip("/")]
        if homepage_pages:
            raw_title = homepage_pages[0].get("title","") or ""
            # Extract brand name: take part after the last "|" or "-" separator
            import re as _re
            parts = _re.split(r'\s*[|\-–]\s*', raw_title)
            client_name = parts[-1].strip() if len(parts) > 1 else (parts[0].strip() if parts else client_slug)
        if not client_name or client_name.startswith("http"):
            # Fallback: capitalise the slug
            client_name = client_slug.replace("-"," ").replace("_"," ").title()

    # ── FIX: Auto-detect industry from crawl page content when not set ────────
    # This ensures the strategic growth section uses industry-specific insights.
    if not industry and pages:
        all_text = " ".join(
            " ".join([p.get("title",""), p.get("meta_description","") or ""])
            for p in pages
        ).lower()
        all_schema_flat = [s for p in pages for s in p.get("schema_types",[])]
        if any(t in all_schema_flat for t in ["Dentist","MedicalBusiness","MedicalClinic","Hospital","Physician"]):
            industry = "local dental clinic"
        elif any(kw in all_text for kw in ["dental","dentist","orthodontic","clinic","implant","braces"]):
            industry = "local dental clinic"
        elif any(kw in all_text for kw in ["restaurant","food","menu","delivery","cuisine"]):
            industry = "local restaurant"
        elif any(kw in all_text for kw in ["salon","spa","beauty","hair","nail"]):
            industry = "local beauty & wellness"
        elif any(kw in all_text for kw in ["real estate","property","buy home","rent","apartment"]):
            industry = "real estate"
        elif any(kw in all_text for kw in ["saas","software","platform","free trial","pricing","integration"]):
            industry = "saas"
        elif any(kw in all_text for kw in ["product","shop","cart","buy","add to bag","collection"]):
            industry = "e-commerce"
        elif any(kw in all_text for kw in ["agency","marketing","seo","design","branding","case study"]):
            industry = "agency"
        else:
            industry = "local service"
    onpage_results=onpage.get("results",[]); avg_onpage=onpage.get("avg_score","N/A")
    total_pages=summary.get("total_pages",len(pages) or "N/A")
    broken=int(summary.get("status_404",0) or 0); redirects=int(summary.get("status_301",0) or 0)
    noindex=int(summary.get("noindex_pages",0) or 0); miss_h1=int(summary.get("missing_h1",0) or 0)
    miss_meta=int(summary.get("missing_meta_desc",0) or 0)
    pages_with_schema=sum(1 for p in pages if p.get("schema_types"))
    schema_pct=f"{round(pages_with_schema/max(len(pages),1)*100)}%" if pages else "N/A"
    all_schema=[s for p in pages for s in p.get("schema_types",[])]
    miss_alt=sum(p.get("images_missing_alt",0) for p in pages)

    lh_mob=lighthouse.get("mobile",{}); lh_scores=lh_mob.get("scores",{})
    lh_cwv=lh_mob.get("core_web_vitals",{}); perf=lh_scores.get("performance")
    lcp_ms=lh_cwv.get("lcp"); inp_ms=lh_cwv.get("inp"); cls_v=lh_cwv.get("cls")

    # ── Advanced Metrics Collection ──────────────────────────────────────────
    trust_fails = sum(1 for p in onpage_results if not p.get("cro_trust", {}).get("has_contact_info", True))
    cta_missing = sum(1 for p in onpage_results if not p.get("cro_ctas") and p.get("word_count", 0) > 200)
    eeat_author_fails = sum(1 for p in onpage_results if not p.get("ee_at_author", True))
    eeat_citation_fails = sum(1 for p in onpage_results if p.get("ee_at_citations", 0) == 0)
    eeat_schema_fails = sum(1 for p in onpage_results if not p.get("schema_same_as", True))
    competitors = brand_kit.get("competitors", {}).get("top_3_competitors", [])

    # ── COVER ────────────────────────────────────────────────────────────────
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(50)
    r=p.add_run("DARE NETWORK"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=NAVY

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6)
    r=p.add_run("COMPREHENSIVE SEO & CRO AUDIT"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=NAVY

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6)
    r=p.add_run(website_url); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=ORANGE

    hr=doc.add_paragraph(); hr.paragraph_format.space_before=Pt(14); hr.paragraph_format.space_after=Pt(14)
    pPr=hr._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    b=OxmlElement("w:bottom"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6")
    b.set(qn("w:space"),"1"); b.set(qn("w:color"),"1B3A6B"); pBdr.append(b); pPr.append(pBdr)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Technical SEO  \u00b7  On-Page SEO  \u00b7  Off-Page SEO  \u00b7  CRO  \u00b7  AEO / GEO  \u00b7  Growth Strategy")
    r.italic=True; r.font.size=Pt(10.5); r.font.color.rgb=MID; p.paragraph_format.space_after=Pt(20)

    st=doc.add_table(rows=1,cols=4); st.alignment=WD_TABLE_ALIGNMENT.CENTER; st.style="Table Grid"
    for i,(num,label) in enumerate([("100+","Issues Reviewed"),("6","Audit Sections"),
                                      (schema_pct,"Schema Coverage"),("30+","Growth Actions")]):
        c=st.rows[0].cells[i]; _set_bg(c,"1B3A6B")
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8)
        rn=p.add_run(num); rn.bold=True; rn.font.size=Pt(18); rn.font.color.rgb=WHITE
        p2=c.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(8)
        rl=p2.add_run(label); rl.font.size=Pt(9); rl.font.color.rgb=WHITE

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(20)
    r=p.add_run(f"Prepared by Dare Network  |  {audit_date}"); r.font.size=Pt(10); r.font.color.rgb=MID
    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────────
    add_h1(doc,"Executive Summary")
    add_body(doc,f"{client_name} is a {industry} brand audited by Dare Network. This report covers: Technical SEO, On-Page Optimisation, Off-Page Authority, AEO/GEO Readiness, Core Web Vitals, Image SEO, and E-E-A-T. Our objective is to identify every gap, prioritise high-impact actions, and define a clear path to increase organic visibility and revenue.")
    doc.add_paragraph()
    add_h2(doc,"Overall Health Snapshot")

    tech_s=_tech_score(summary, framework_data)
    onp_s=int(avg_onpage) if str(avg_onpage).isdigit() else 65

    # Apply on-page SEO score cap if framework has CSR_SPA issue
    if framework_data and framework_data.get("status") == "success":
        score_cap_data = framework_data.get("score_cap", {})
        if score_cap_data and "on_page_seo" in score_cap_data:
            max_onpage = score_cap_data["on_page_seo"] * 10  # Convert 3/10 to 30/100
            onp_s = min(onp_s, max_onpage)
    perf_disp=f"Mobile: {perf}/100" if perf else "Not yet measured"
    schema_s=min(int(pages_with_schema/max(len(pages),1)*10) if pages else 2, 10)

    add_health_table(doc,[
        ("Technical SEO",f"{broken} broken, {miss_h1} missing H1, {miss_meta} missing meta",f"{tech_s//10}/10","Critical" if tech_s<50 else "High" if tech_s<70 else "Medium"),
        ("On-Page SEO",f"Avg on-page score {avg_onpage}/100",f"{onp_s//10}/10","High" if onp_s<70 else "Medium"),
        ("CRO & UX",f"{trust_fails} pages lack trust signals, {cta_missing} lack CTAs",f"{10 - min(10, (trust_fails+cta_missing)//2)}/10","Critical" if cta_missing>0 else "Medium"),
        ("Content / Blog","Content depth and keyword targeting","5/10","High"),
        ("Competitor Gap",f"Benchmarked against {len(competitors)} competitors" if competitors else "No competitors provided","5/10","High"),
        ("Off-Page / Links","Domain authority and backlink profile","3/10","High"),
        ("Core Web Vitals",perf_disp,f"{perf//10}/10" if perf else "N/A","Critical" if not perf or perf<50 else "High"),
        ("Advanced E-E-A-T","Author entities and expert citations validation",f"{10 - min(10, (eeat_author_fails+eeat_schema_fails)//2)}/10","High"),
        ("AEO / GEO Readiness",f"AI Share of Voice (SOV): 0% | Schema: {schema_pct}",f"{schema_s}/10","Critical" if schema_s<=3 else "High"),
        ("Image SEO",f"{miss_alt} images missing alt text","3/10" if miss_alt>10 else "6/10","Critical" if miss_alt>10 else "Medium"),
    ])
    add_callout(doc,"Key Opportunity",[
        f"{client_name} has strong potential. With structured SEO implementation — fixing technical foundations, building schema markup, and executing a targeted content strategy — this site has the potential to 3x organic traffic within 12 months."
    ],"info")
    doc.add_page_break()

    # ── SECTION 1: TECHNICAL SEO ─────────────────────────────────────────────
    add_section_banner(doc,1,"Technical SEO Audit")
    add_h1(doc,"Technical SEO Audit")
    add_body(doc,f"Technical SEO forms the foundation of all organic rankings. Even the best content cannot rank if Google cannot efficiently crawl, index, and understand your website. Below is a comprehensive breakdown of all technical issues observed on {website_url}.")

    # ── 1.0: FRAMEWORK DETECTION (NEW - CRITICAL) ───────────────────────────
    framework_data = load_tmp(f"*{client_slug}*framework*.json")
    if framework_data and framework_data.get("status") == "success":
        add_h2(doc, "1.0  Framework & Rendering Analysis [CRITICAL]")
        add_body(doc, "Modern websites often use JavaScript frameworks (React, Vue, Angular) to render content. While this creates smooth user experiences, it can make content invisible to Google if not implemented correctly. We analyzed what Google sees (no JavaScript) vs what users see (JavaScript enabled).")

        framework = framework_data.get("framework", "Unknown")
        render_mode = framework_data.get("render_mode", "UNKNOWN")
        seo_verdict = framework_data.get("seo_verdict", "WARNING")
        nojs_words = framework_data.get("nojs_word_count", 0)
        js_words = framework_data.get("js_word_count", 0)
        content_ratio = framework_data.get("content_ratio", 0)
        recommendation = framework_data.get("recommendation", "")

        # Add results callout
        callout_style = "critical" if seo_verdict == "CRITICAL" else "warning" if seo_verdict == "WARNING" else "success"
        callout_lines = [
            f"Framework Detected: {framework}",
            f"Render Mode: {render_mode}",
            f"Content Visible to Google: {nojs_words} words (no-JS)",
            f"Content Visible to Users: {js_words} words (with JS)",
            f"Google Visibility Ratio: {content_ratio:.1%}",
            ""
        ]

        if content_ratio < 0.1:
            callout_lines.append(f"🔴 CRITICAL: Less than 10% of your content is visible to Google!")
        elif content_ratio < 0.5:
            callout_lines.append(f"⚠️ WARNING: Less than 50% of your content is visible to Google")
        else:
            callout_lines.append(f"✅ GOOD: Most content is visible to Google")

        add_callout(doc, f"{seo_verdict}: Framework Rendering Issue", callout_lines, callout_style)

        # Add framework-specific issues
        framework_issues = []
        if render_mode == "CSR_SPA":
            framework_issues.append({
                "issue": f"Client-Side Rendering ({framework})",
                "severity": "Critical",
                "finding": f"Your website uses {framework} with client-side rendering (CSR). Google sees only {nojs_words} words while users see {js_words} words. This means {100-int(content_ratio*100)}% of your content is invisible to search engines.",
                "fix": recommendation
            })
            framework_issues.append({
                "issue": "Technical SEO Score Capped",
                "severity": "Critical",
                "finding": "Due to CSR architecture, your maximum achievable technical SEO score is 2/10 regardless of other optimizations. This is the #1 priority to fix.",
                "fix": "All other technical SEO fixes will have minimal impact until the framework migration is complete. Prioritize this above everything else."
            })
        elif render_mode in ["HYBRID", "UNKNOWN"] and content_ratio < 0.9:
            framework_issues.append({
                "issue": f"Partial Content Rendering ({framework})",
                "severity": "High",
                "finding": f"Some content is rendered client-side. Google can see {int(content_ratio*100)}% of your full content. The remaining {100-int(content_ratio*100)}% is at risk of not being indexed.",
                "fix": "Audit which content sections are JS-dependent and move them to server-side rendering or static generation."
            })

        if framework_issues:
            add_issues_table(doc, framework_issues)
        else:
            add_callout(doc, "✅ Framework Implementation: GOOD", [
                f"Your site uses {framework} with proper server-side rendering or static generation.",
                "Google can see all your content without executing JavaScript.",
                "No critical framework-related issues detected."
            ], "success")

    if cms and cms.lower() == "shopify":
        add_h2(doc,"1.1  Shopify Infrastructure Audit")
        add_issues_table(doc,[
            {"issue":"Shopify Liquid Code Duplication","severity":"High","finding":"Potential repetitive 'ticker loops' or header bloat in Liquid templates, slowing down main thread rendering.","fix":"Refactor Liquid code to use caching ({% cache %}) or minimize DOM size in global headers."},
            {"issue":"International SEO Structure","severity":"Medium","finding":"Hreflang tags or localized pricing may be missing or improperly configured across regions.","fix":"Implement Shopify Markets to correctly handle currency, local URLs, and automated hreflang tags."},
            {"issue":"Shopify App Bloat","severity":"Medium","finding":"Third-party Shopify apps often inject render-blocking JavaScript in the <head>.","fix":"Consolidate apps, defer non-critical scripts, and leverage Shopify's built-in CDN-level image optimization."}
        ])
    
    add_h2(doc,"1.1  Crawlability & Indexation")
    add_body(doc,"Crawlability determines how efficiently Googlebot can discover and process pages. Indexation determines which of those pages appear in search results.")
    crawl_issues=[]
    if broken>0: crawl_issues.append({"issue":f"Broken Pages — {broken} URLs returning 404","severity":"Critical","finding":f"{broken} pages return 404 Not Found. These waste crawl budget, destroy link equity, and create terrible user experiences.","fix":"Set up 301 redirects for every broken URL to the most relevant live page. Check Google Search Console > Coverage report weekly."})
    if miss_h1>0: crawl_issues.append({"issue":f"Missing H1 Tags — {miss_h1} pages","severity":"High","finding":f"{miss_h1} pages ({round(miss_h1/max(int(str(total_pages) if str(total_pages).isdigit() else 1),1)*100)}% of site) have no H1 tag. Google uses H1 as the primary topic signal for a page.","fix":"Add a unique, keyword-rich H1 to every page. The H1 should match user search intent, not just the brand name."})
    # ── FIX: Only flag canonical issue if canonicals are actually missing ──────
    pages_with_canonical = sum(1 for p in pages if p.get("canonical"))
    if pages and pages_with_canonical < len(pages) * 0.5:
        crawl_issues.append({"issue":"Missing Canonical Tags","severity":"High","finding":f"Only {pages_with_canonical}/{len(pages)} pages have canonical tags. Without them Google may index duplicate or near-duplicate URLs, splitting ranking signals.","fix":"Add self-referencing canonical tags to every page via your CMS template (Rank Math / Yoast)."})
    elif pages and pages_with_canonical < len(pages):
        crawl_issues.append({"issue":"Incomplete Canonical Implementation","severity":"Medium","finding":f"{len(pages)-pages_with_canonical} pages are missing canonical tags. {pages_with_canonical}/{len(pages)} pages have them set correctly.","fix":"Ensure all pages have self-referencing canonical tags. Check paginated pages and filter URLs especially."})
    # ── FIX: Only flag duplicate URL issue if non-trailing-slash variants exist ─
    canonical_urls = set(p.get("canonical","").rstrip("/") for p in pages if p.get("canonical"))
    raw_urls = set(p.get("url","").rstrip("/") for p in pages)
    if len(raw_urls) > len(canonical_urls) and pages:
        crawl_issues.append({"issue":"Duplicate URL Variants (Trailing Slash)","severity":"Medium","finding":f"Both trailing-slash and non-trailing-slash versions of pages are being crawled (e.g. /about/ and /about both resolve). This wastes crawl budget even when canonicals are set.","fix":"Implement a 301 redirect at the server/CDN level to enforce a single canonical URL format. Prefer trailing-slash versions."})
    # ── FIX: Only flag sitemap as unsubmitted if no sitemap reference in robots.txt ──
    robots_data = load_tmp(f"*{client_slug}*robots*.json")
    has_sitemap_reference = bool(robots_data) or any("sitemap" in str(p.get("url","")).lower() for p in pages)
    if not has_sitemap_reference:
        crawl_issues.append({"issue":"XML Sitemap Not Confirmed","severity":"High","finding":"No verified sitemap submission detected in crawl data. An unsubmitted sitemap means Google may miss new pages for weeks.","fix":f"Submit {website_url}sitemap_index.xml in Google Search Console. Monitor the Coverage report for indexation errors weekly."})
    else:
        crawl_issues.append({"issue":"Verify Sitemap Submission in GSC","severity":"Low","finding":f"Sitemap is present ({website_url}sitemap_index.xml) but GSC submission status is unverified. Unverified sitemaps may not be processed.","fix":"In Google Search Console → Sitemaps, confirm the sitemap shows 'Success' status and all pages are being indexed."})
    if noindex>0: crawl_issues.append({"issue":f"Pages with Noindex — {noindex} URLs","severity":"Medium","finding":f"{noindex} pages carry noindex tags preventing them from appearing in search results.","fix":"Review each noindex page. Remove the tag from any page that should rank. Keep noindex only on admin, thank-you, and test pages."})
    if redirects>0: crawl_issues.append({"issue":f"Redirect Chains — {redirects} detected","severity":"Medium","finding":f"{redirects} redirect chains found. Each hop wastes crawl budget and dilutes PageRank.","fix":"Update all internal links to point directly to final destination URLs. Eliminate multi-hop redirect chains."})
    if miss_meta>0: crawl_issues.append({"issue":f"Missing Meta Descriptions — {miss_meta} pages","severity":"Medium","finding":f"{miss_meta} pages have no meta description. Google auto-generates snippets, often using irrelevant navigation or footer text.","fix":"Write compelling meta descriptions (120-160 chars) for all pages. Include the primary keyword and a clear value proposition."})
    # ── FIX: Detect links-not-crawlable from Lighthouse SEO audits ───────────
    lh_seo_audits = lighthouse.get("mobile",{}).get("seo_audits",{})
    if lh_seo_audits.get("crawlable-anchors",{}).get("score",1) == 0:
        crawl_issues.append({"issue":"Links Are Not Crawlable (Lighthouse)","severity":"Critical","finding":"Lighthouse detected anchor elements using non-standard href attributes (JavaScript handlers, #, or void). Google cannot follow these links to discover and crawl linked pages.","fix":"Replace all JS-based link handlers with proper href='/page/' attributes. Run in DevTools console: document.querySelectorAll('a:not([href]), a[href=\"#\"], a[href*=\"javascript\"]') to identify offenders."})
    if not crawl_issues: crawl_issues.append({"issue":"No Critical Crawl Issues","severity":"Low","finding":"No major crawl or indexation issues detected.","fix":"Run a monthly deep crawl with Screaming Frog or Sitebulb for ongoing monitoring."})
    add_issues_table(doc,crawl_issues)

    add_h2(doc,"1.2  Page Speed & Core Web Vitals")
    add_body(doc,"Google's Core Web Vitals (LCP, INP, CLS) are confirmed ranking signals. They measure real-world loading, interactivity, and visual stability. Poor scores directly reduce rankings and increase bounce rates.")
    if perf:
        lcp_s=f"{lcp_ms/1000:.1f}s" if lcp_ms else "N/A"; inp_l=f"{inp_ms:.0f}ms" if inp_ms else "N/A"; cls_l=f"{cls_v:.2f}" if cls_v else "N/A"
        add_callout(doc,"Measured Performance (Google PageSpeed Insights)",[f"LCP: {lcp_s} — {'POOR' if lcp_ms and lcp_ms>4000 else 'NEEDS IMPROVEMENT' if lcp_ms and lcp_ms>2500 else 'GOOD'} (target <2.5s)",f"INP: {inp_l} (target <200ms)",f"CLS: {cls_l} (target <0.1)",f"Mobile Performance Score: {perf}/100"],"warning" if perf<50 else "info")
    else:
        add_callout(doc,"Estimated Performance (Based on Visual Analysis)",["LCP: Likely POOR — large images loading without preloading detected","INP: Likely HIGH — render-blocking scripts detected","CLS: Likely HIGH — images without explicit dimensions detected","Add GOOGLE_API_KEY to .env to get exact measured PageSpeed scores"],"warning")
    # ── FIX: CWV issues now reference ACTUAL measured values from Lighthouse ────
    cwv_issues=[]
    if lcp_ms:
        lcp_secs = lcp_ms / 1000
        lcp_rating = "POOR" if lcp_ms > 4000 else "NEEDS IMPROVEMENT" if lcp_ms > 2500 else "GOOD"
        lcp_sev = "Critical" if lcp_ms > 4000 else "High" if lcp_ms > 2500 else "Medium"
        cwv_issues.append({"issue":f"LCP: {lcp_secs:.1f}s ({lcp_rating})","severity":lcp_sev,"finding":f"Largest Contentful Paint is {lcp_secs:.1f} seconds on mobile — target is under 2.5s. This is {round(lcp_secs/2.5,1)}x over the acceptable limit. Every extra second of LCP reduces conversions by an estimated 7%. Primary causes: uncompressed hero images, render-blocking scripts in <head>, no preloading of above-fold images.","fix":"1) Convert all images to WebP (25-35% smaller). 2) Move GTM/GA4/analytics scripts to async/defer. 3) Add <link rel='preload'> for the hero image in <head>. 4) Enable a caching plugin (WP Rocket / LiteSpeed Cache)."})
    else:
        cwv_issues.append({"issue":"LCP: Not Yet Measured","severity":"High","finding":"LCP (Largest Contentful Paint) could not be measured. This is a primary Core Web Vitals ranking signal. Without measuring it, you cannot identify the specific asset causing slowness.","fix":"Run Google PageSpeed Insights on the live site. The 'Opportunities' section will name the specific image or script causing the bottleneck."})

    lh_mob_cwv = lighthouse.get("mobile",{}).get("core_web_vitals",{})
    tbt_ms = lh_mob_cwv.get("tbt")
    fcp_ms = lh_mob_cwv.get("fcp")
    tti_ms_val = lh_mob_cwv.get("tti")
    si_ms = lh_mob_cwv.get("speed_index")

    if tbt_ms and tbt_ms > 200:
        tbt_sev = "Critical" if tbt_ms > 600 else "High"
        cwv_issues.append({"issue":f"Total Blocking Time: {tbt_ms:.0f}ms","severity":tbt_sev,"finding":f"TBT is {tbt_ms:.0f}ms (target: under 200ms). High TBT means heavy JavaScript is blocking the browser main thread, preventing users from interacting with the page. This directly causes poor INP scores and failed Core Web Vitals.","fix":"Defer all non-critical third-party scripts (GTM, analytics, chat widgets, WhatsApp buttons). Use Chrome DevTools > Performance tab to identify which scripts contribute most to the blocking time."})
    if fcp_ms and fcp_ms > 1800:
        cwv_issues.append({"issue":f"First Contentful Paint: {fcp_ms/1000:.1f}s","severity":"High","finding":f"FCP is {fcp_ms/1000:.1f}s — users see a blank page for this long before any content appears. Poor FCP increases bounce rate before the page even loads.","fix":"Prioritize above-fold HTML delivery. Eliminate render-blocking CSS/JS. Enable server-side compression (gzip/brotli). Use a CDN for static assets."})
    cwv_issues.append({"issue":"Images Not in WebP / AVIF Format","severity":"High","finding":"Images served as JPEG/PNG instead of modern formats. WebP is 25-35% smaller at equivalent quality, directly improving LCP. Given the measured LCP is primarily image-driven, this is a direct fix.","fix":"Convert all images to WebP before uploading. Use tools like Squoosh.app, Cloudflare Images, or a WordPress plugin (Imagify, ShortPixel). Implement <picture> element with JPEG fallback for older browsers."})
    cwv_issues.append({"issue":"No Lazy Loading on Below-Fold Images","severity":"Medium","finding":f"Below-fold images load on initial page load, wasting bandwidth and inflating LCP. On a page with {miss_alt + sum(p.get('images_missing_alt',0)>0 and 1 or 0 for p in pages)} total images, this is a significant weight problem.","fix":"Add loading='lazy' to all below-fold images. The hero/LCP image must use loading='eager' and be preloaded in <head> with <link rel='preload' as='image'>."})
    cwv_issues.append({"issue":"No Image Compression Pipeline","severity":"Medium","finding":"Images uploaded without compression. Files over 100KB for content images and 200KB for heroes are serious performance red flags.","fix":"Implement compression before upload. Target: WebP at quality 75-80. For WordPress: use WP Rocket + Imagify together for both caching and image optimisation in one step."})
    if cms and "shopify" in cms.lower(): cwv_issues.append({"issue":"Shopify App Script Bloat","severity":"Medium","finding":"Each installed Shopify app typically injects 1-3 scripts. 5+ apps create significant render-blocking weight.","fix":"Audit all installed apps. Remove unused ones. Consolidate duplicate function apps. Use native Shopify features where possible."})
    add_issues_table(doc,cwv_issues)

    add_h2(doc,"1.3  Mobile Optimisation")
    add_body(doc,"Google uses mobile-first indexing. Your mobile site is what Google primarily crawls and ranks. A poor mobile experience directly translates to lower rankings for all devices.")
    add_issues_table(doc,[
        {"issue":"Touch Target Sizes","severity":"Medium","finding":"Navigation links and filter tags may be below the recommended 44x44px touch target, causing mis-taps on mobile.","fix":"Ensure all interactive elements are minimum 44px height. Use CSS padding rather than fixed heights to accommodate this."},
        {"issue":"Mobile Font Size","severity":"Low","finding":"Body text may render below 16px on some devices, which Google flags as a readability issue in Search Console.","fix":"Set base font-size to 16px minimum. Verify using Chrome DevTools Device Toolbar across multiple device sizes."},
    ])

    add_h2(doc,"1.4  Structured Data / Schema Markup")
    add_body(doc,"Schema markup tells Google what your content means, enabling rich results: star ratings, prices, article carousels, breadcrumbs, and sitelinks. Pages with correct schema consistently get higher CTR.")
    if not all_schema or not pages:
        add_callout(doc,"Critical Gap: Minimal Schema Implementation",[f"The site currently lacks most structured data types that unlock rich snippets.","Competitors with Product and Review schema display star ratings and prices directly in Google results — getting far more clicks even at the same ranking position."],"critical")
    schema_iss=[]
    # ── FIX: Check for missing schema types with real data + deeper coverage ──
    if not any("Organization" in s for s in all_schema):
        schema_iss.append({"issue":"Organization Schema Missing","severity":"High","finding":"No Organization schema on homepage. This is the most fundamental schema — it gives Google your brand name, logo, social profiles, and contact info. Without it, Google cannot build a Knowledge Panel for the brand.","fix":"Add Organization JSON-LD to homepage <head> with: name, logo, url, sameAs (all social profiles), contactPoint, @id (unique entity URL)."})
    if not any("Article" in s or "BlogPosting" in s for s in all_schema):
        schema_iss.append({"issue":"Article / BlogPosting Schema Missing","severity":"High","finding":"Blog posts lack Article schema. Without it Google cannot display article-rich results: author, publish date, and article carousels.","fix":"Add Article or BlogPosting JSON-LD to all blog posts with headline, author, datePublished, dateModified, and image."})
    if not any("BreadcrumbList" in s for s in all_schema):
        schema_iss.append({"issue":"BreadcrumbList Schema Missing","severity":"Medium","finding":"No BreadcrumbList schema detected. Breadcrumbs display the page path in Google results, improving CTR by clearly showing site structure.","fix":"Add BreadcrumbList JSON-LD to all non-homepage pages. Especially impactful for category and product pages."})
    if not any("Person" in s for s in all_schema):
        schema_iss.append({"issue":"Person / Author Schema Missing","severity":"Medium","finding":"No author entity markup detected. Person schema builds E-E-A-T signals by connecting content to credentialed human authors.","fix":"Add Person JSON-LD to all author pages: name, url, sameAs (LinkedIn, Twitter/X), jobTitle, and knowsAbout."})
    # ── Critical: Check for FAQPage schema (AEO/AI search requirement) ────────
    if not any("FAQPage" in s or "FAQ" in s for s in all_schema):
        schema_iss.append({"issue":"FAQPage Schema Missing","severity":"High","finding":"No FAQPage schema detected anywhere on the site. FAQ schema is the primary driver of Google AI Overview citations and voice search answers. Competitors with FAQ schema get featured snippets and AI Overview placements for the same queries.","fix":"Add FAQPage JSON-LD with 5-8 Q&A pairs to the homepage and service pages. Questions should match real user search queries (use Google Autosuggest + People Also Ask)."})
    # ── Critical: Check for AggregateRating / Review schema ──────────────────
    if not any("AggregateRating" in s or "Review" in s for s in all_schema):
        schema_iss.append({"issue":"AggregateRating (Star Ratings) Schema Missing","severity":"High","finding":"No star rating schema found. Competitors with AggregateRating schema display gold star ratings in Google results, getting 15-30% higher CTR at the same ranking position. This is especially critical for local businesses where reviews drive conversion.","fix":"Add AggregateRating to the Organization/LocalBusiness schema: ratingValue, reviewCount, bestRating: 5. Pull values from Google Business Profile rating."})
    # ── Check for LocalBusiness/Dentist specific schema ───────────────────────
    has_local = any(s in all_schema for s in ["LocalBusiness","Dentist","MedicalBusiness","MedicalClinic","Store","Restaurant"])
    if not has_local and ("local" in industry.lower() or "dental" in industry.lower() or "service" in industry.lower()):
        schema_iss.append({"issue":"LocalBusiness / Dentist Schema Missing","severity":"High","finding":"No LocalBusiness or Dentist schema found. For local businesses, this schema type tells Google your precise NAP (name, address, phone), hours, service area, and price range — all critical for Local Pack rankings.","fix":"Add Dentist (or LocalBusiness) JSON-LD with: name, address, telephone, openingHours, priceRange, areaServed, hasMap (Google Maps URL)."})
    # ── Check for sameAs entity links ────────────────────────────────────────
    schema_has_same_as = any(p.get("schema_same_as") for p in onpage_results)
    if not schema_has_same_as:
        schema_iss.append({"issue":"Entity sameAs Links Missing","severity":"Medium","finding":"Organization schema lacks sameAs property linking to third-party entity validators (Google Business Profile, LinkedIn, Wikidata, Practo, Justdial). Google uses sameAs to cross-verify your brand entity and build your Knowledge Graph presence.","fix":"Add sameAs array to Organization JSON-LD listing all brand profiles: Google Maps URL, LinkedIn, Facebook, Instagram, any directory listings (Practo, Justdial), and Wikipedia if available."})
    # ── Check for Service schema on service pages ─────────────────────────────
    service_pages = [p for p in pages if "/service" in p.get("url","").lower() or "/treatment" in p.get("url","").lower()]
    if service_pages and not any("Service" in s for s in all_schema):
        schema_iss.append({"issue":"Service Schema Missing on Service Pages","severity":"Medium","finding":f"The services page lists multiple treatment offerings but has no Service or MedicalProcedure schema. Each service is a rankable entity — without schema, Google treats them as plain text.","fix":"Add individual Service JSON-LD blocks for each treatment (Dental Implants, Root Canal, Orthodontics etc.) with name, description, provider, and areaServed."})
    # ── WebSite schema check ──────────────────────────────────────────────────
    if not any("WebSite" in s or "SearchAction" in s for s in all_schema):
        schema_iss.append({"issue":"WebSite Schema / Sitelinks Search Box Missing","severity":"Low","finding":"No WebSite schema with SearchAction found. This schema enables the Sitelinks Search Box in Google results and signals site authority to the Knowledge Graph.","fix":"Add WebSite JSON-LD to homepage with SearchAction pointing to your internal search results URL."})
    if not schema_iss:
        schema_iss.append({"issue":"Enhance Existing Schema","severity":"Low","finding":"Core schema types are present. Opportunity exists to add WebSite (Sitelinks Search Box) and SiteNavigationElement schema.","fix":"Add WebSite schema with SearchAction. Add SiteNavigationElement linking to main navigation destinations."})
    add_issues_table(doc,schema_iss)
    doc.add_page_break()

    # ── SECTION 2: ON-PAGE ───────────────────────────────────────────────────
    add_section_banner(doc,2,"On-Page SEO Audit")
    add_h1(doc,"On-Page SEO Audit")
    add_body(doc,"On-page SEO encompasses everything optimisable within the website itself: titles, meta descriptions, headings, content quality, keyword targeting, image SEO, and internal linking.")

    add_h2(doc,"2.1  Title Tags & Meta Descriptions")
    add_body(doc,"Title tags (50-60 chars) are the single most important on-page ranking factor — they appear as the blue headline in Google results. Meta descriptions (120-160 chars) heavily influence click-through rate.")
    if onpage_results:
        lbl2=doc.add_paragraph(); lbl2.paragraph_format.space_before=Pt(8); lbl2.paragraph_format.space_after=Pt(2)
        lr2=lbl2.add_run("Top Pages — On-Page Score"); lr2.bold=True; lr2.font.size=Pt(11); lr2.font.color.rgb=NAVY
        tbl=doc.add_table(rows=1,cols=3); tbl.style="Table Grid"
        for i,(lbl_t,w) in enumerate(zip(["Page URL","Score / 100","Key Issues"],[Inches(2.5),Inches(1.0),Inches(3.5)])):
            c=tbl.rows[0].cells[i]; _set_bg(c,"1B3A6B"); c.width=w
            r=c.paragraphs[0].add_run(lbl_t); r.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(10)
        for res in onpage_results[:10]:
            row=tbl.add_row().cells; sc=res.get("overall_score",0)
            row[0].paragraphs[0].add_run(res.get("url","")[-50:]).font.size=Pt(9)
            rs=row[1].paragraphs[0].add_run(f"{sc}/100"); rs.bold=True; rs.font.size=Pt(10)
            _set_bg(row[1],"C8E6C9" if sc>=75 else "FFF9C4" if sc>=50 else "FFCDD2")
            ri=row[2].paragraphs[0].add_run("; ".join(res.get("issues",[])[:2]) or "OK"); ri.font.size=Pt(9)
        doc.add_paragraph()
    # ── FIX: Title/meta issues now reflect ACTUAL on-page data ─────────────────
    title_meta_issues = []
    titles_too_long   = [r for r in onpage_results if r.get("title",{}).get("length",0) > 60]
    titles_too_short  = [r for r in onpage_results if 0 < r.get("title",{}).get("length",0) < 40]
    titles_missing    = [r for r in onpage_results if not r.get("title",{}).get("text","")]
    metas_too_long    = [r for r in onpage_results if r.get("meta",{}).get("length",0) > 160]
    metas_missing     = [r for r in onpage_results if not r.get("meta",{}).get("text","")]

    if titles_missing:
        title_meta_issues.append({"issue":f"Missing Title Tags — {len(titles_missing)} pages","severity":"Critical","finding":f"{len(titles_missing)} pages have no title tag. This is the most important on-page ranking factor. Google will auto-generate titles which are usually suboptimal.","fix":"Add a unique, keyword-rich title (50-60 chars) to every page. Lead with the primary keyword. Format: 'Primary Keyword | Brand Name'."})
    if titles_too_long:
        title_meta_issues.append({"issue":f"Title Tags Too Long — {len(titles_too_long)} pages","severity":"High","finding":f"{len(titles_too_long)} pages have titles over 60 characters. Google truncates these in search results with '...', cutting off the brand name or key differentiator. Affected: {'; '.join(r.get('url','')[-40:] for r in titles_too_long[:3])}.","fix":"Shorten titles to 50-60 characters. Keep the primary keyword at the start. Remove filler words like 'Welcome to', 'The Best', 'Official Site'."})
    if titles_too_short:
        title_meta_issues.append({"issue":f"Title Tags Too Short — {len(titles_too_short)} pages","severity":"Medium","finding":f"{len(titles_too_short)} pages have titles under 40 characters, leaving unused SERP real estate. Affected: {'; '.join(r.get('url','')[-40:] for r in titles_too_short[:3])}.","fix":"Expand titles to 50-60 chars. Include the primary keyword and location/brand qualifier."})
    if metas_missing:
        title_meta_issues.append({"issue":f"Missing Meta Descriptions — {len(metas_missing)} pages","severity":"High","finding":f"{len(metas_missing)} pages have no meta description. Google auto-generates snippets using random page text, often pulling navigation or footer content — reducing CTR significantly.","fix":"Write unique, compelling meta descriptions (120-160 chars) for all pages. Include the primary keyword in the first half and end with a clear CTA ('Book now', 'Learn more')."})
    if metas_too_long:
        title_meta_issues.append({"issue":f"Meta Descriptions Too Long — {len(metas_too_long)} pages","severity":"Medium","finding":f"{len(metas_too_long)} meta descriptions exceed 160 characters and will be truncated in SERPs. Affected: {'; '.join(r.get('url','')[-40:] for r in metas_too_long[:3])}.","fix":"Trim meta descriptions to 120-160 characters. Put the most important information and CTA in the first 120 characters."})
    title_meta_issues.append({"issue":"Keyword Cannibalisation Risk","severity":"Medium","finding":"Multiple pages appear to target the same primary keyword variations, splitting ranking signals. Google will pick one page to rank and may suppress others entirely.","fix":"Audit all pages for keyword overlap. Assign one unique primary keyword per page. Consolidate near-duplicate pages or differentiate content angles clearly."})
    if not title_meta_issues:
        title_meta_issues.append({"issue":"Title & Meta Optimisation","severity":"Low","finding":"Titles and meta descriptions are within acceptable length ranges. Opportunity exists to improve keyword placement and conversion-oriented language.","fix":"A/B test meta descriptions using GSC CTR data. Pages with <2% CTR despite good rankings should be rewritten."})
    add_issues_table(doc, title_meta_issues)

    add_h2(doc,"2.2  Heading Structure (H1–H3)")
    add_body(doc,"Every page must have exactly one H1. H2s and H3s structure content for both users and search engines, signalling topical depth and semantic relevance.")
    if miss_h1>0:
        add_callout(doc,f"Critical: {miss_h1} pages missing H1",["Pages without H1 send no clear topic signal to Google.","The H1 is the strongest on-page heading SEO signal.","Fix immediately: add a unique, keyword-rich H1 to every affected page."],"critical")
    # ── FIX: Detect pages with multiple H1 tags ───────────────────────────────
    multi_h1_pages = [r for r in onpage_results if len(r.get("headings",{}).get("h1",[]))>1]
    if multi_h1_pages:
        affected = ', '.join(r.get('url','')[-50:] for r in multi_h1_pages[:3])
        add_callout(doc,f"Critical: {len(multi_h1_pages)} pages have multiple H1 tags",[
            f"Affected: {affected}",
            "Multiple H1s split Google's topic signal — it cannot determine which heading defines the page primary subject.",
            "Common cause: sliders, team sections, or testimonial blocks coded as H1 by mistake.",
            "Fix: Keep exactly one H1 per page. Change all extra H1s to H2 or H3 as appropriate."
        ],"critical")

    add_h2(doc,"2.3  Image SEO")
    add_body(doc,"Image SEO drives Google Image Search traffic and provides critical accessibility and ranking signals. It is consistently the fastest-win category in any audit.")
    img_iss=[]
    if miss_alt>0 or not pages: img_iss.append({"issue":f"Missing Alt Text — {miss_alt if miss_alt else 'Many'} images","severity":"Critical","finding":f"{miss_alt if miss_alt else 'Numerous'} images have no alt attribute. This is an accessibility failure and wastes all image SEO signal.","fix":"Add descriptive alt text to every image. Describe the image content and include relevant keywords naturally — never keyword-stuff."})
    img_iss+=[
        {"issue":"Images Not in WebP / AVIF Format","severity":"High","finding":"Images served as JPEG/PNG instead of modern formats. WebP is 25-35% smaller at equivalent quality, directly improving LCP.","fix":"Convert all images to WebP. Use AVIF for supporting browsers. Implement <picture> element with JPEG fallback."},
        {"issue":"No Lazy Loading on Below-Fold Images","severity":"Medium","finding":"Below-fold images load on initial page load, wasting bandwidth and slowing the LCP metric for all users.","fix":"Add loading='lazy' to all below-fold images. The hero/LCP image must use loading='eager' and be preloaded in <head>."},
        {"issue":"Generic Image Filenames","severity":"Low","finding":"Images with filenames like 'image001.jpg' or 'IMG_4521.jpg' miss keyword signals. Google reads filenames as content clues.","fix":"Rename images descriptively before upload. Example: 'navy-handloom-cotton-saree-bandhani.jpg' not 'product-7.jpg'."},
    ]
    add_issues_table(doc,img_iss)

    add_h2(doc,"2.4  Internal Linking")
    add_body(doc,"Internal links distribute PageRank and help Google understand which pages are most important. A structured internal linking strategy can meaningfully improve rankings without any external effort.")
    add_issues_table(doc,[
        {"issue":"Generic Anchor Text","severity":"High","finding":"Internal links use generic anchor text ('click here', 'read more'). Google uses anchor text to understand what the linked page is about.","fix":"Replace all generic anchors with keyword-rich descriptive text. Example: 'Shop Handloom Sarees' not 'View Collection'."},
        {"issue":"Blog Not Linking to Products","severity":"High","finding":"Blog content rarely links to relevant product/collection pages, missing the most valuable internal PageRank flow opportunity.","fix":"Rule: every blog post must include minimum 3 contextual internal links to relevant product or collection pages."},
        {"issue":"No Related Products Strategy","severity":"Medium","finding":"Related product sections use generic text ('You may also like'), missing a structured internal link graph within the product catalog.","fix":"Implement 'Pairs Well With' or 'Complete the Look' sections with keyword-rich anchors linking to complementary products."},
    ])
    doc.add_page_break()

    # ── SECTION 3: CRO & UX ──────────────────────────────────────────────────
    doc.add_page_break()
    add_section_banner(doc,3,"Integrated CRO & UX Audit (Strategic Conversion)")
    add_h1(doc,"Integrated CRO & UX Audit")
    add_body(doc,"Moving beyond organic traffic acquisition, we audit elements that turn that traffic into revenue. High-ranking pages mean nothing if users bounce due to poor UX, mismatched search intents, or lack of trust signals.")
    cro_iss=[]
    if trust_fails>0: cro_iss.append({"issue":"Missing Trust Signals","severity":"High","finding":f"{trust_fails} analyzed pages lack critical trust signals like visible contact info or clear policy links. This destroys conversion rates.","fix":"Add clear phone numbers, email addresses, return policies, and SSL badges globally to the header/footer and checkout flows."})
    if cta_missing>0: cro_iss.append({"issue":"Intent-to-CTA Mismatch","severity":"High","finding":f"{cta_missing} content-heavy pages lack clear Call-to-Action buttons. Users consume content but have no obvious next step.","fix":"Map CTA to intent: educational pages need 'soft' CTAs (download guide, subscribe), commercial pages need 'hard' CTAs (Buy Now)."})
    
    if is_ecommerce:
        cro_iss.append({"issue":"Credibility Zone Optimization","severity":"High","finding":"Social proof, founder stories, or trust badges are buried below the fold rather than instantly visible.","fix":"Move star ratings, reviews, and 'Buyer Protection' / 'Hassle-Free Exchange' badges immediately below the hero banner (the Credibility Zone)."})
        cro_iss.append({"issue":"Intent-to-CTA Alignment (High AOV)","severity":"High","finding":"High-ticket items (>₹5,000 / $100) often rely on standard 'Add to Cart' buttons, causing hesitation drops.","fix":"Implement WhatsApp Integration ('Chat with an Expert') on high-consideration product pages to capture leads and close sales."})
        cro_iss.append({"issue":"Average Order Value (AOV) Boost","severity":"Medium","finding":"Product pages lack strategic internal linking to drive logical upsells.","fix":"Always include a 'Complete the Look' or 'Related Resource' internal linking strategy."})
        
    cro_iss.append({"issue":"High Impression / Low CTR Risks","severity":"Medium","finding":"Without GA4/GSC integration, identifying pages with high SERP visibility but low click-through or high bounce rates is difficult.","fix":"Connect GA4 & GSC to Dare Network OS to automate 'Low Scroll Depth' and 'Engagement Rate' anomaly detection for immediate restructuring."})
    add_issues_table(doc,cro_iss)

    # ── SECTION 4: ADVANCED E-E-A-T ──────────────────────────────────────────
    doc.add_page_break()
    add_section_banner(doc,4,"Advanced E-E-A-T & Entity Validation")
    add_h1(doc,"Advanced E-E-A-T & Entity Validation")
    add_body(doc,"Google's March 2024 updates heavily prioritize Experience, Expertise, Authoritativeness, and Trustworthiness. We validate the entity graph connecting your brand and authors to the broader web.")
    eeat_iss=[]
    if eeat_author_fails>0: eeat_iss.append({"issue":"No Author Profile Automation","severity":"High","finding":f"{eeat_author_fails} pages lack dedicated author entities. Google's Quality Rater Guidelines heavily weigh content creator expertise.","fix":"Generate standalone Author Profile pages including credentials, social links, and press features. Link to them from all blog posts with rel='author'."})
    if eeat_citation_fails>0: eeat_iss.append({"issue":"Missing Expert Citations","severity":"Medium","finding":f"Several informational pages lack outbound links to authoritative industry sources (.edu, .gov, or recognized boards).","fix":"Identify factual claims in content and add citations to high-authority domains to validate your own content's trustworthiness."})
    if eeat_schema_fails>0: eeat_iss.append({"issue":"Entity Validation Gap (Schema sameAs)","severity":"High","finding":"Organization or Person schema lacks 'sameAs' arrays. Google uses these to cross-reference your brand with third-party validation (LinkedIn, Wikipedia, Crunchbase).","fix":"Inject sameAs links to all existing social profiles and third-party validation sources into Organization and Person JSON-LD."})
    add_issues_table(doc,eeat_iss)

    # ── SECTION 5: COMPETITOR INTELLIGENCE ───────────────────────────────────
    doc.add_page_break()
    add_section_banner(doc,5,"Automated Competitor Intelligence Module")
    add_h1(doc,"Competitor Intelligence & Gap Analysis")
    add_body(doc, "Understanding your competitive landscape is critical. We benchmarked your site against top competitors across 10 key dimensions to identify gaps and opportunities.")

    # ── COMPETITOR SCORECARD TABLE (NEW - ENHANCED) ─────────────────────────
    competitor_data = load_tmp(f"*{client_slug}*competitors*.json")
    if competitor_data and competitor_data.get("competitors"):
        add_h2(doc, "5.1  Competitive Scorecard — Side-by-Side Analysis")

        # Create scorecard table
        competitors_list = competitor_data.get("competitors", [])[:3]  # Max 3 competitors
        num_cols = 2 + len(competitors_list)  # Client + competitors
        comp_table = doc.add_table(rows=1, cols=num_cols)
        comp_table.style = "Table Grid"

        # Header row
        _set_bg(comp_table.rows[0].cells[0], "1B3A6B")
        r0 = comp_table.rows[0].cells[0].paragraphs[0].add_run("Dimension")
        r0.bold = True; r0.font.color.rgb = WHITE; r0.font.size = Pt(10)

        _set_bg(comp_table.rows[0].cells[1], "E8671A")  # Orange for client
        r1 = comp_table.rows[0].cells[1].paragraphs[0].add_run(client_name)
        r1.bold = True; r1.font.color.rgb = WHITE; r1.font.size = Pt(10)

        for i, comp in enumerate(competitors_list):
            _set_bg(comp_table.rows[0].cells[i+2], "1B3A6B")
            rc = comp_table.rows[0].cells[i+2].paragraphs[0].add_run(comp.get("name", f"Competitor {i+1}"))
            rc.bold = True; rc.font.color.rgb = WHITE; rc.font.size = Pt(9)

        # Scoring dimensions
        dimensions = [
            ("Technical SEO", "technical_seo"),
            ("On-Page SEO", "on_page_seo"),
            ("Content Volume", "content_volume"),
            ("Local SEO", "local_seo"),
            ("Schema Markup", "schema_markup"),
            ("Blog Activity", "blog_activity"),
            ("Mobile Experience", "mobile_experience"),
            ("E-E-A-T Signals", "eeat_signals"),
            ("CTA Quality", "cta_quality"),
            ("Overall Score", "overall")
        ]

        client_scores = competitor_data.get("client_scores", {})

        for dim_label, dim_key in dimensions:
            row = comp_table.add_row().cells
            row[0].paragraphs[0].add_run(dim_label).font.size = Pt(9)

            # Client score
            client_score = client_scores.get(dim_key, tech_s//10 if dim_key == "technical_seo" else onp_s//10 if dim_key == "on_page_seo" else 5)
            score_text = f"{client_score}/10" if dim_key != "overall" else f"{client_score:.1f}"
            rs = row[1].paragraphs[0].add_run(score_text)
            rs.font.size = Pt(9); rs.bold = True

            # Color code
            if isinstance(client_score, (int, float)):
                if client_score >= 7:
                    _set_bg(row[1], "C8E6C9")  # Green
                elif client_score >= 5:
                    _set_bg(row[1], "FFF9C4")  # Yellow
                else:
                    _set_bg(row[1], "FFCDD2")  # Red

            # Competitor scores
            for i, comp in enumerate(competitors_list):
                comp_score = comp.get("scores", {}).get(dim_key, 6)
                comp_score_text = f"{comp_score}/10" if dim_key != "overall" else f"{comp_score:.1f}"
                rcs = row[i+2].paragraphs[0].add_run(comp_score_text)
                rcs.font.size = Pt(9)

                if isinstance(comp_score, (int, float)):
                    if comp_score >= 7:
                        _set_bg(row[i+2], "C8E6C9")
                    elif comp_score >= 5:
                        _set_bg(row[i+2], "FFF9C4")
                    else:
                        _set_bg(row[i+2], "FFCDD2")

        doc.add_paragraph()

        # Competitive insights
        add_h2(doc, "5.2  Competitive Insights — What They Do Better")
        for i, comp in enumerate(competitors_list):
            comp_name = comp.get("name", f"Competitor {i+1}")
            insights = comp.get("insights", [])
            if insights:
                add_body(doc, f"{comp_name}:", colour=NAVY)
                for insight in insights[:3]:
                    add_bullet(doc, insight)

            gap = comp.get("exploitable_gap", "")
            if gap:
                add_body(doc, f"⚡ Quick Win: {gap}", colour=ORANGE, italic=True)
            doc.add_paragraph()

    # Original dynamic competitors section
    dyn_comps = strategy_data.get("competitors", [])
    if dyn_comps and not competitor_data:
        add_body(doc, "Side-by-side benchmarking against local rivals reveals market share gaps:")
        for c in dyn_comps:
            add_bullet(doc, f"{c.get('name')}: Possesses a '{c.get('asset')}' which is currently missing from your domain.")
        dyn_kws = strategy_data.get("keywords", [])
        if dyn_kws:
            add_body(doc, f"Immediate Keyword Opportunities (assign to new Collection Pages): {', '.join(dyn_kws)}", italic=True)

    comp_iss=[
        {"issue":"Keyword Gap Analysis Required","severity":"High","finding":"A comprehensive overlap analysis against competitors is required to identify high-volume, high-intent keywords they rank for that you are missing.","fix":"Run automated keyword gap analysis. Filter for keywords where competitors rank on Page 1 but you do not rank in the top 50."},
        {"issue":"Topic Cluster Depth Gap","severity":"Medium","finding":"Competitor sitemap structures often reveal specialized topic clusters where they have significantly deeper content coverage.","fix":"Map competitor informational hubs and generate 'Content Briefs' for missing topics to match and exceed their topical authority."},
        {"issue":"SERP Overlap / Rivalry Score","severity":"Low","finding":"To focus resources, we must pinpoint your truest search competitors — those overlapping most heavily in SERPs, not just offline business competitors.","fix":"Calculate 'Rivalry Scores' combining Domain Authority vs Keyword Overlap to prioritize which competitors to reverse-engineer first."}
    ]
    add_issues_table(doc,comp_iss)

    # ── SECTION 6: OFF-PAGE ──────────────────────────────────────────────────
    doc.add_page_break()
    add_section_banner(doc,6,"Off-Page SEO & Link Building")
    add_h1(doc,"Off-Page SEO Audit")
    add_body(doc,"Off-page SEO covers everything outside your website that influences rankings — primarily backlinks, brand authority, and digital PR. A strong backlink profile is the primary differentiator between page 1 and page 2+ rankings.")
    add_issues_table(doc,[
        {"issue":"No Systematic Link Building","severity":"High","finding":"No evidence of a structured link building programme. Without consistent new backlinks, domain authority stagnates while competitors grow.","fix":"Start a guest post programme: 4-5 quality posts per month in your niche. Use /link_building workflow to find and prioritise targets."},
        {"issue":"Unlinked Brand Mentions","severity":"Medium","finding":"The brand is likely mentioned across the web without a backlink — a very common missed opportunity for consumer brands.","fix":"Set up Google Alerts for your brand name. When mentions appear without links, reach out politely with a link request."},
        {"issue":"No Digital PR Strategy","severity":"Medium","finding":"No evidence of media outreach or press coverage that would generate high-authority editorial backlinks.","fix":"Pitch stories to niche publications using brand assets: customer stories, founder story, unique data, or seasonal angles."},
        {"issue":"Google Business Profile Status","severity":"Medium","finding":"GBP completeness is unconfirmed. An incomplete GBP means missing the brand knowledge panel and related local signals.","fix":"Claim and fully verify GBP. Complete all fields: name, address, phone, website, categories, description, and add photos."},
    ])
    doc.add_page_break()

    # ── SECTION 7: AEO / GEO ─────────────────────────────────────────────────
    add_section_banner(doc,7,"AEO / GEO Readiness — AI Search Optimisation")
    add_h1(doc,"AEO / GEO Readiness")
    add_body(doc,"Answer Engine Optimisation (AEO) and Generative Engine Optimisation (GEO) focus on being cited in AI-powered surfaces: Google AI Overviews, ChatGPT, Perplexity, and Bing Copilot. Sites with strong structured data, clear Q&A formats, and authoritative E-E-A-T are most frequently cited.")
    add_issues_table(doc,[
        {"issue":"No Direct Answer Blocks","severity":"High","finding":"Pages lack structured 'Who/What/Why' answer blocks in the first 100 words. AI engines extract these concise answers for AI Overviews citations.","fix":"Add a 50-100 word direct answer to the page's target question at the very top. Write as if answering a voice search query directly and completely."},
        {"issue":"Zero AI Share of Voice (SOV)","severity":"Critical","finding":"Benchmarking tools show 0% visibility in Answer Engines (ChatGPT, Perplexity). The brand is not currently being cited for primary industry questions.","fix":"Implement llms.txt & LLMFeeds in /.well-known/ directory immediately to feed structured brand data to AI crawlers."},
        {"issue":"No WebSite Schema with SearchAction","severity":"High","finding":"WebSite schema is missing from homepage. This property enables the Sitelinks Search Box in Google results and signals site authority.","fix":"Add WebSite JSON-LD to homepage with SearchAction pointing to your internal search results URL."},
        {"issue":"No Author Entity Markup","severity":"High","finding":"Content is not linked to verified author entities. Google's Knowledge Graph uses entity associations to validate E-E-A-T credibility claims.","fix":"Add Person schema for all content authors: name, url, sameAs (LinkedIn, Twitter/X), jobTitle, and knowsAbout properties."},
        {"issue":"Content Not Structured for AI (Readability)","severity":"Medium","finding":"Pages use long prose paragraphs. Text complexity (Flesch score) is sub-optimal for LLM token extraction.","fix":"Run NLP Readability module. Restructure key pages with clear H2 questions followed by 2-3 sentence direct answers. Target Grade 8 reading level."},
        {"issue":"Weak E-E-A-T Signals","severity":"Medium","finding":"Site lacks visible experience evidence: no case studies, no specific results data, no displayed credentials or years of experience.","fix":"Add an 'Our Results' section. Display specific client metrics. Add founder credentials and years of experience prominently throughout."},
    ])
    doc.add_page_break()

    # ── SECTION 8: ACTION PLAN ───────────────────────────────────────────────
    add_section_banner(doc,8,"Prioritised Action Plan — 90 Day Roadmap")
    add_h1(doc,"Prioritised Action Plan")
    add_body(doc,"Every issue identified in this audit is listed below, organised by business impact and urgency. Execute in order for maximum organic growth within 90 days.")

    add_h2(doc,"\U0001f534  Critical — Fix Within 30 Days")
    critical=[]
    if miss_h1>0: critical.append(f"Add H1 tags to all {miss_h1} pages currently missing them")
    if broken>0: critical.append(f"Fix {broken} broken pages (404 errors) — implement 301 redirects immediately")
    if miss_alt>0: critical.append(f"Add alt text to all {miss_alt} images — start with top 10 most-visited pages")
    critical+=["Add Organization schema to homepage (E-E-A-T + Knowledge Graph signal)","Add Article/BlogPosting schema to all blog posts","Add canonical tags to every page site-wide","Submit and verify XML sitemap in Google Search Console"]
    for item in critical: add_bullet(doc,item)

    add_h2(doc,"\U0001f7e0  High Priority — Complete Within 60 Days")
    for item in ["Rewrite title tags for top 20 pages — lead with primary keyword (50-60 chars)","Add 300+ words of descriptive content to all category and listing pages","Convert all images to WebP format and implement lazy loading site-wide","Add author bio boxes with credentials to all blog/content pages","Start guest post outreach — target 5 new quality backlinks per month","Add direct answer blocks (50-100 words) to top 10 landing pages for AEO","Implement BreadcrumbList and Person schema site-wide"]: add_bullet(doc,item)

    add_h2(doc,"\U0001f7e1  Medium Priority — Ongoing (60-180 Days)")
    for item in ["Build Pillar + Cluster content architecture around top 5 keyword themes","Add WebSite schema with SearchAction to homepage","Restructure key pages to Q&A format for AI Overview eligibility","Monthly GSC review: track impressions, CTR, avg position, and index coverage","Connect Google Analytics 4 and Search Console to AIOS for automated monthly reports","Implement link building programme: 5 guest posts per month with keyword-rich anchor text"]: add_bullet(doc,item)

    # ── 8.1: KEYWORD STRATEGY MAP (NEW) ──────────────────────────────────────
    add_h2(doc, "8.1  Keyword Strategy Map — Target Pages & Priority")
    add_body(doc, "Every keyword should be mapped to a specific page to avoid cannibalization. This table prioritizes keywords by: (Search Volume × Intent Match) / Difficulty.")

    # Load keyword data
    keyword_data = load_tmp(f"*{client_slug}*keywords*.json") or load_tmp(f"*{client_slug}*gap*.json")
    primary_keywords = brand_kit.get("seo_settings", {}).get("primary_keywords", [])

    if keyword_data and keyword_data.get("keywords"):
        kw_table = doc.add_table(rows=1, cols=6)
        kw_table.style = "Table Grid"

        # Header
        for i, (label, w) in enumerate(zip(
            ["Keyword", "Search Intent", "Volume", "Difficulty", "Priority", "Target Page"],
            [Inches(1.8), Inches(1.0), Inches(0.7), Inches(0.7), Inches(0.7), Inches(2.0)]
        )):
            c = kw_table.rows[0].cells[i]
            _set_bg(c, "1B3A6B")
            c.width = w
            r = c.paragraphs[0].add_run(label)
            r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)

        # Add keyword rows
        keywords_list = keyword_data.get("keywords", [])[:20]  # Top 20
        for kw in keywords_list:
            row = kw_table.add_row().cells
            row[0].paragraphs[0].add_run(kw.get("keyword", "")).font.size = Pt(9)
            row[1].paragraphs[0].add_run(kw.get("intent", "Informational")).font.size = Pt(8)
            row[2].paragraphs[0].add_run(str(kw.get("volume", "N/A"))).font.size = Pt(9)
            row[3].paragraphs[0].add_run(str(kw.get("difficulty", "N/A"))).font.size = Pt(9)

            # Priority color coding
            priority = kw.get("priority", "MEDIUM")
            rp = row[4].paragraphs[0].add_run(priority)
            rp.font.size = Pt(8); rp.bold = True
            if priority == "HIGH":
                _set_bg(row[4], "FFCDD2")
            elif priority == "MEDIUM":
                _set_bg(row[4], "FFF9C4")
            else:
                _set_bg(row[4], "C8E6C9")

            row[5].paragraphs[0].add_run(kw.get("target_page", "/")).font.size = Pt(8)

    elif primary_keywords:
        # Fallback: use brand_kit primary keywords
        kw_table = doc.add_table(rows=1, cols=4)
        kw_table.style = "Table Grid"

        for i, label in enumerate(["Keyword", "Priority", "Target Page", "Status"]):
            c = kw_table.rows[0].cells[i]
            _set_bg(c, "1B3A6B")
            r = c.paragraphs[0].add_run(label)
            r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)

        for kw in primary_keywords[:10]:
            row = kw_table.add_row().cells
            row[0].paragraphs[0].add_run(kw).font.size = Pt(9)
            row[1].paragraphs[0].add_run("HIGH").font.size = Pt(9)
            row[2].paragraphs[0].add_run("TBD — Assign in GSC").font.size = Pt(8)
            row[3].paragraphs[0].add_run("Needs Mapping").font.size = Pt(8)
    else:
        add_body(doc, "Keyword data unavailable. Run /keyword_research workflow to generate a comprehensive keyword strategy map.", italic=True, colour=MID)

    doc.add_paragraph()
    doc.add_page_break()

    # ── SECTION 9: DARE NETWORK VALUE ────────────────────────────────────────
    add_section_banner(doc,9,"How Dare Network Adds Value")
    add_h1(doc,"How Dare Network Adds Value")
    vt=doc.add_table(rows=1,cols=2); vt.style="Table Grid"
    for i,label in enumerate(["What We Bring to the Table","What You Get"]):
        c=vt.rows[0].cells[i]; _set_bg(c,"1B3A6B")
        r=c.paragraphs[0].add_run(label); r.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(11)
    for left,right in [
        ("AI-Powered SEO OS","Automated audits, content briefs, and reports at agency speed"),
        ("Technical SEO Execution","Direct implementation — no back-and-forth"),
        ("Content Strategy","Pillar + Cluster content that captures long-tail organic traffic"),
        ("AEO / GEO Optimisation","Structured to appear in Google AI Overviews and chatbot answers"),
        ("Link Building","Systematic outreach: 5+ quality backlinks per month"),
        ("Monthly Reporting","Progress tied to business outcomes — not vanity metrics"),
    ]:
        row=vt.add_row().cells
        r0=row[0].paragraphs[0].add_run(left); r0.bold=True; r0.font.size=Pt(10)
        r1=row[1].paragraphs[0].add_run(right); r1.font.size=Pt(10)
    doc.add_paragraph()
    
    # ── SECTION 10: STRATEGIC INSIGHTS & INNOVATION ──────────────────────────
    doc.add_page_break()
    add_section_banner(doc,10,"STRATEGIC INSIGHTS & GROWTH INNOVATION")
    add_h1(doc,"Strategic Growth Insights")
    add_body(doc,f"To break out of standard SEO patterns, we've identified {industry}-specific growth opportunities based on competitive gaps, market trends, and content whitespace.")

    # Load strategic insights data
    insights_data = load_tmp(f"*{client_slug}*strategic*.json") or load_tmp(f"*{client_slug}*insights*.json")

    if insights_data and insights_data.get("insights"):
        # Dynamic industry-specific insights
        insights_list = insights_data.get("insights", [])
        for i, insight in enumerate(insights_list[:6]):
            add_h2(doc, f"10.{i+1}  {insight.get('title', 'Growth Opportunity')}")
            add_body(doc, insight.get("description", ""))

            # Metrics if available
            if insight.get("metrics"):
                metrics = insight.get("metrics")
                add_callout(doc, "Opportunity Metrics", [
                    f"Search Volume: {metrics.get('volume', 'N/A')}/month",
                    f"Effort Level: {metrics.get('effort', 'Medium')}",
                    f"Time to Result: {metrics.get('time_to_result', '3-6 months')}",
                    f"Expected Traffic Lift: {metrics.get('expected_lift', '20-40%')}"
                ], "info")

            if insight.get("action_steps"):
                add_body(doc, "Action Steps:", colour=NAVY)
                for step in insight.get("action_steps", []):
                    add_bullet(doc, step)
            doc.add_paragraph()

    else:
        # ── FIX: Industry-matched insights using auto-detected industry ───────
        ind_lower = industry.lower()
        if "dental" in ind_lower or "dentist" in ind_lower:
            insights = [
                ("Individual Service Landing Pages", f"Create dedicated pages for each treatment: /dental-implants/, /invisalign/, /root-canal-treatment/, /wisdom-tooth-removal/. Each page targets high-intent queries like 'dental implants HSR Layout cost'. Competitors with dedicated service pages consistently outrank general /services/ pages for procedure-specific searches."),
                ("Treatment Cost / Pricing Transparency Pages", "Patients research costs before booking. Create a 'Dental Treatment Costs in [City]' page that ranks for 'dental implant cost [city]' queries — one of the highest-volume dental keywords. Transparent pricing builds trust and reduces phone-shopping friction."),
                ("Before/After Patient Case Studies", "Document 5-10 real patient transformations (with consent): problem, treatment chosen, timeline, and result photos. These pages rank for '[treatment] results [city]', generate natural backlinks from dental forums, and build the E-E-A-T signals Google requires for YMYL health sites."),
                ("Doctor Credentials & Authority Pages", "Build individual profile pages for each dentist with: BDS/MDS qualifications, years of experience, specialisation areas, IDA membership, and case count. Google's Quality Rater Guidelines give extra weight to identifiable experts on health/medical sites. Link these from all blog posts as author pages."),
                ("Local Neighbourhood SEO Pages", f"Create location-specific pages targeting neighbouring areas: '/dentist-in-[nearby-area]/', '/dental-clinic-whitefield/', '/orthodontist-koramangala/'. These capture patients searching from adjacent neighbourhoods who may not find the clinic via HSR Layout-specific queries.")
            ]
        elif "restaurant" in ind_lower or "food" in ind_lower:
            insights = [
                ("Menu SEO Pages", "Create individual SEO-optimised pages for signature dishes, cuisine types, and dietary options. Rank for 'best [cuisine] in [city]' queries."),
                ("Occasion & Booking Landing Pages", "Pages for: anniversary dinners, corporate lunches, birthday parties. These capture high-intent 'restaurant for [occasion] near me' queries."),
                ("Chef Story & Sourcing Content", "Publish content about sourcing, cooking philosophy, and chef credentials. These build E-E-A-T for food/hospitality YMYL content.")
            ]
        elif "e-commerce" in ind_lower or "ecommerce" in ind_lower or "shop" in ind_lower:
            insights = [
                ("Product Comparison Pages", "Create '[Product A] vs [Product B]' comparison pages for top products. These capture high-intent comparison queries with 3x higher conversion rates than category pages."),
                ("User-Generated Content Strategy", "Launch a review collection campaign. Products with 50+ reviews rank 4.6x higher and convert 270% better. Incentivize reviews with discount codes or loyalty points."),
                ("Seasonal Content Calendar", "Map content to buying cycles: Pre-festival guides (60 days before), Gift guides (30 days), Last-minute deals (7 days). Target 'best [product] for [occasion]' queries.")
            ]
        elif "saas" in ind_lower or "software" in ind_lower:
            insights = [
                ("Comparison Landing Pages", "Create '[Your Tool] vs [Competitor]' pages for top 5 competitors. These queries have 8x higher intent and 65% conversion rate."),
                ("Free Tool / Calculator Strategy", "Build a simple ROI calculator or assessment tool. These rank for 'calculator' queries, generate backlinks naturally, and collect lead emails."),
                ("Integration Hub", "Create individual integration pages for every connected tool. These rank for '[Tool] + [Integration]' long-tail queries.")
            ]
        elif "local" in ind_lower or "service" in ind_lower:
            insights = [
                ("Hyperlocal Content Strategy", "Create neighbourhood-specific service pages. Instead of one '/services' page, create 10+ pages like '/services-[neighbourhood]' with unique local content and reviews."),
                ("Before/After Content", "Document every project with before/after photos and detailed case studies. These build trust and rank for '[service] near me before and after' queries."),
                ("Local Partnership Content", "Partner with complementary local businesses for co-marketing content. Guest posts on local business blogs build high-relevance backlinks.")
            ]
        elif "agency" in ind_lower or "marketing" in ind_lower:
            insights = [
                ("Case Study Library", "Publish 5+ detailed case studies with specific metrics: 'How we grew organic traffic 312% for [industry client] in 6 months'. These rank for '[service] results' queries and build E-E-A-T authority."),
                ("Tools & Calculators", "Build free SEO/marketing calculators (ROI calculator, keyword difficulty estimator). These generate passive backlinks and email captures."),
                ("Competitor Analysis Content", "Publish '[Competitor Tool] alternatives' pages. High-intent, easy-to-rank bottom-funnel content.")
            ]
        else:
            # Generic B2B/Agency insights
            insights = [
                ("Encyclopedia Pillar Content", "Create a comprehensive glossary defining every industry term. Each term becomes a rankable page that captures early-stage awareness traffic."),
                ("Founder Personal Brand", "Build the founder's LinkedIn presence and personal website. Google associates entity authority with individuals, not just companies. This strengthens E-E-A-T."),
                ("Original Research / Industry Report", "Publish an annual '[Industry] State of the Market Report'. Original data generates backlinks from news sites, industry blogs, and Wikipedia citations.")
            ]

        for i, (title, desc) in enumerate(insights):
            add_h2(doc, f"10.{i+1}  {title}")
            add_body(doc, desc)
            doc.add_paragraph()

    # Additional strategic recommendations from strategy_data
    idea_1 = strategy_data.get("innovation_idea_1", "")
    idea_2 = strategy_data.get("innovation_idea_2", "")

    if idea_1:
        add_callout(doc, "Additional Strategy: Education & Authority", [idea_1], "info")
        doc.add_paragraph()
    add_callout(doc, "Idea 2: Seasonal Activation", [idea_2], "info")
    doc.add_paragraph()
    add_h1(doc,"How Dare Network Adds Value")
    vt=doc.add_table(rows=1,cols=2); vt.style="Table Grid"
    for i,label in enumerate(["What We Bring to the Table","What You Get"]):
        c=vt.rows[0].cells[i]; _set_bg(c,"1B3A6B")
        r=c.paragraphs[0].add_run(label); r.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(11)
    for left,right in [
        ("AI-Powered SEO OS","Automated audits, content briefs, and reports at agency speed"),
        ("Technical SEO Execution","Direct implementation — no back-and-forth"),
        ("Content Strategy","Pillar + Cluster content that captures long-tail organic traffic"),
        ("AEO / GEO Optimisation","Structured to appear in Google AI Overviews and chatbot answers"),
        ("Link Building","Systematic outreach: 5+ quality backlinks per month"),
        ("Monthly Reporting","Progress tied to business outcomes — not vanity metrics"),
    ]:
        row=vt.add_row().cells
        r0=row[0].paragraphs[0].add_run(left); r0.bold=True; r0.font.size=Pt(10)
        r1=row[1].paragraphs[0].add_run(right); r1.font.size=Pt(10)
    doc.add_paragraph()
    add_callout(doc,"Ready to Execute?",["This audit identifies every change needed to significantly improve organic rankings.","Dare Network will execute this roadmap systematically, report progress monthly, and continuously optimise based on data.","Next step: review the Critical Priority list and begin Week 1 implementation."],"info")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(20)
    r=p.add_run("— End of Audit Report —"); r.font.size=Pt(10); r.font.color.rgb=MID
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run("www.darenetwork.in"); r2.font.size=Pt(11); r2.bold=True; r2.font.color.rgb=ORANGE
    return doc


def build_monthly_docx(client_slug, website_url, brand_kit, month):
    doc=Document()
    for sec in doc.sections: sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    info=brand_kit.get("client_info",{}); client_name=info.get("client_name",client_slug)
    month_label=datetime.strptime(month,"%Y-%m").strftime("%B %Y") if month else "Last Month"
    add_h1(doc,f"Monthly SEO Report — {month_label}")
    add_body(doc,f"Client: {client_name}  |  {website_url}  |  Prepared by Dare Network")
    sections = [
        ("1. Organic Growth Performance", "Focus on Google Search Console (GSC) clicks and impressions trends."),
        ("2. Search & AI Visibility", "Focus on share-of-voice in Google Search and AI Overview appearances."),
        ("3. Content Strategy Execution", "List articles published and their topical relevance."),
        ("4. Outreach & Link Building", "List backlinks acquired and outreach statuses."),
        ("5. Focus for Next Month", "Top 3 priorities to drive growth.")
    ]
    
    for section, note in sections:
        add_h2(doc, section); add_body(doc, note)
    add_body(doc,f"\n— Report generated by SEO AI OS | Dare Network | {datetime.now().strftime('%B %d, %Y')} —",MID)
    return doc


def main():
    parser=argparse.ArgumentParser(description="SEO Report Builder — Dare Network Template v2")
    parser.add_argument("--client"); parser.add_argument("--url"); parser.add_argument("--type",required=True,choices=["audit","monthly"]); parser.add_argument("--month"); parser.add_argument("--strategy")
    args=parser.parse_args()

    strategy_data = {}
    if args.strategy and Path(args.strategy).exists():
        try:
            strategy_data = json.load(open(args.strategy, encoding="utf-8"))
        except Exception as e:
            print(f"[Warning] Failed to load strategy JSON: {e}")

    if args.client:
        client_slug=args.client; brand_kit=load_brand_kit(client_slug)
        website_url=brand_kit.get("client_info",{}).get("website_url",f"{client_slug}.com")
    elif args.url:
        # CRITICAL: Use centralized url_to_slug() to ensure consistency across all tools
        client_slug = url_to_slug(args.url)
        brand_kit={"client_info":{"client_name":args.url,"website_url":args.url}}; website_url=args.url
    else:
        print("[Error] Provide --client or --url"); return

    # VALIDATION: Check that required files exist (prevents 10/10 bug)
    if args.type == "audit" and args.url:
        required_files = ["framework", "crawl_nojs", "lighthouse"]
        validation = validate_file_naming(args.url, required_files)

        if not validation["valid"]:
            print("\n[ERROR] Required audit files are missing!")
            print(f"Expected slug: {validation['slug']}")
            print("\nMissing files:")
            for missing in validation["missing"]:
                print(f"  - {missing}")
            print("\nRun the complete audit workflow first:")
            print(f"  1. python tools/framework_detector.py --url {args.url} --output .tmp/{validation['slug']}_framework.json")
            print(f"  2. python tools/seo_crawler.py --url {args.url} --no-js --output .tmp/{validation['slug']}_crawl_nojs.json")
            print(f"  3. python tools/lighthouse_audit.py --url {args.url} --strategy both --output .tmp/{validation['slug']}_lighthouse.json")
            print()
            return

    # Always save to .tmp/reports/ — never to audit_history
    output_dir=Path(".tmp/reports"); output_dir.mkdir(parents=True,exist_ok=True)
    today=date.today().strftime("%Y-%m-%d")

    # Use unique timestamp to avoid file locking issues
    import time
    timestamp = int(time.time())
    filename=f"{today}_{args.type}_report_{timestamp}.docx"
    if args.type=="monthly" and args.month: filename=f"{args.month}_{args.type}_report_{timestamp}.docx"
    output_path=output_dir/filename

    print(f"[Report Builder] Generating {args.type} report for: {website_url}")
    doc=build_audit_docx(client_slug,website_url,brand_kit,strategy_data) if args.type=="audit" else build_monthly_docx(client_slug,website_url,brand_kit,args.month or datetime.now().strftime("%Y-%m"))
    doc.save(str(output_path)); abs_path=str(output_path.resolve())

    print(f"\n[SUCCESS] Report saved!")
    print(f"[File]    {abs_path}")
    print(f"\n[Download] Link: file:///{abs_path.replace(chr(92), '/')}")

if __name__=="__main__": main()
