#!/usr/bin/env python3
"""Generate AGENCY_PLAYBOOK.docx from AGENCY_PLAYBOOK.md"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x3A, 0x6B)
ORANGE = RGBColor(0xE8, 0x67, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x33, 0x33, 0x33)

def build_playbook_docx():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    md_path = Path(__file__).parent.parent / "AGENCY_PLAYBOOK.md"
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False

    for line in lines:
        line = line.rstrip()

        # Track code blocks
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        
        # Render code block lines as monospace
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(line)
            r.font.size = Pt(9)
            r.font.name = "Consolas"
            r.font.color.rgb = DARK
            continue

        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = NAVY
            continue

        # H2
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = NAVY
            continue

        # H3
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line[4:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = NAVY
            continue

        # H4
        if line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(line[5:])
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = DARK
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph()
            continue

        # Bullet points
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            _add_formatted_text(p, text)
            continue

        # Numbered items
        m = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            _add_formatted_text(p, text)
            continue

        # Table rows (skip for now - tables rendered as text)
        if line.strip().startswith("|"):
            text = line.strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            r.font.name = "Consolas"
            r.font.color.rgb = DARK
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_formatted_text(p, text)

    out = Path(__file__).parent.parent / "AGENCY_PLAYBOOK.docx"
    doc.save(str(out))
    print(f"[OK] AGENCY_PLAYBOOK.docx generated: {out}")


def _add_formatted_text(paragraph, text):
    """Add text with **bold** and *italic* formatting."""
    # Split on bold markers
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        r = paragraph.add_run(part)
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
        if i % 2 == 1:
            r.bold = True


if __name__ == "__main__":
    build_playbook_docx()
